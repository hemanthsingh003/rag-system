"""RAG Engine - Core RAG logic for document ingestion and querying"""

import os
# Suppress verbose outputs before imports
os.environ['TOKENIZERS_PARALLELISM'] = 'false'
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'true'

import warnings
warnings.filterwarnings('ignore')

import sys
import time
import threading
import subprocess
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

# Suppress logging
logging.getLogger('transformers').setLevel(logging.ERROR)
logging.getLogger('paddle').setLevel(logging.ERROR)

# Check for dependencies
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, VectorParams, PointStruct
    QDRANT_AVAILABLE = True
except ImportError:
    QDRANT_AVAILABLE = False

try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False

try:
    import ftfy
    FTFY_AVAILABLE = True
except ImportError:
    FTFY_AVAILABLE = False

from config import config


import uuid

@dataclass
class Document:
    """Represents a processed document"""
    id: str
    title: str
    text: str
    summary: str
    sections: List['Section']


@dataclass
class Section:
    """Represents a document section"""
    id: str
    title: str
    text: str
    summary: str
    doc_id: str
    doc_title: str
    doc_summary: str


@dataclass
class Chunk:
    """Represents a text chunk"""
    id: str
    text: str
    section_id: str
    doc_id: str
    page: Optional[int] = None


class OllamaManager:
    """Manages Ollama service and model lifecycle"""
    
    def __init__(self, model: str = None, llm_timeout: int = None, ollama_timeout: int = None):
        self.model = model or config.OLLAMA_MODEL
        self.llm_timeout = llm_timeout or config.OLLAMA_TIMEOUT
        self.ollama_timeout = ollama_timeout or config.OLLAMA_STOP_TIMEOUT
        self.llm_unload_timer = None
        self.ollama_stop_timer = None
        self.last_activity = time.time()
    
    def is_service_running(self) -> bool:
        result = subprocess.run(
            ["pgrep", "-f", "ollama"],
            capture_output=True
        )
        return result.returncode == 0
    
    def start_service(self):
        """Start Ollama in background"""
        if self.is_service_running():
            return
        
        subprocess.Popen(
            ["ollama", "serve"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        
        for _ in range(10):
            time.sleep(1)
            if self.is_service_running():
                return
        return False
    
    def is_model_loaded(self) -> bool:
        result = subprocess.run(
            ["ollama", "list"],
            capture_output=True,
            text=True
        )
        return self.model in result.stdout
    
    def generate(self, prompt: str, system_prompt: str = None) -> str | None:
        """Generate response with LLM"""
        self.start_service()
        self._cancel_timers()
        self.last_activity = time.time()
        
        full_prompt = prompt
        if system_prompt:
            full_prompt = f"System: {system_prompt}\n\nUser: {prompt}"
        
        result = subprocess.run(
            ["ollama", "run", self.model],
            input=full_prompt,
            capture_output=True,
            text=True,
            timeout=180
        )
        
        if result.returncode != 0:
            print(f"LLM Error: {result.stderr}", file=sys.stderr)
            return None
        
        self._schedule_unload()
        return result.stdout.strip()
    
    def unload_model(self):
        """Unload model from RAM"""
        if self.is_model_loaded():
            subprocess.run(
                ["ollama", "stop", self.model],
                capture_output=True
            )
    
    def stop_service(self):
        """Stop Ollama service completely"""
        subprocess.run(["pkill", "-f", "ollama"], capture_output=True)
    
    def _cancel_timers(self):
        """Cancel pending timers"""
        if self.llm_unload_timer:
            self.llm_unload_timer.cancel()
        if self.ollama_stop_timer:
            self.ollama_stop_timer.cancel()
    
    def _schedule_unload(self):
        """Schedule LLM unload after inactivity"""
        self._cancel_timers()
        
        self.llm_unload_timer = threading.Timer(self.llm_timeout, self._on_llm_idle)
        self.llm_unload_timer.daemon = True
        self.llm_unload_timer.start()
        
        self.ollama_stop_timer = threading.Timer(self.ollama_timeout, self._on_ollama_idle)
        self.ollama_stop_timer.daemon = True
        self.ollama_stop_timer.start()
    
    def _on_llm_idle(self):
        """Called when LLM has been idle"""
        self.unload_model()
    
    def _on_ollama_idle(self):
        """Called when Ollama has been idle too long"""
        self.stop_service()


class EmbeddingManager:
    """Manages embedding model for text vectorization"""
    
    _instance = None
    _model = None
    
    @classmethod
    def get_model(cls, persist: bool = None):
        """Get or load embedding model"""
        if persist is None:
            persist = config.EMBEDDING_PERSIST
        
        if cls._model is None and SENTENCE_TRANSFORMERS_AVAILABLE:
            cls._model = SentenceTransformer(config.EMBEDDING_MODEL)
        
        return cls._model
    
    @classmethod
    def embed_texts(cls, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for texts"""
        model = cls.get_model()
        embeddings = model.encode(texts, convert_to_numpy=True)
        return embeddings.tolist()
    
    @classmethod
    def embed_query(cls, query: str) -> List[float]:
        """Generate embedding for query"""
        model = cls.get_model()
        embedding = model.encode(query, convert_to_numpy=True)
        return embedding.tolist()
    
    @classmethod
    def unload(cls):
        """Unload embedding model"""
        if cls._model is not None:
            del cls._model
            cls._model = None


class DocumentProcessor:
    """Processes documents: OCR, cleaning, chunking"""
    
    def __init__(self):
        self.ocr = None
    
    def process_file(self, file_path: str) -> Document:
        """Process a document file"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            text = self._process_pdf(path)
        elif ext in ['.txt', '.md']:
            text = self._process_text(path)
        elif ext == '.docx':
            text = self._process_docx(path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            text = self._process_image(path)
        elif ext in ['.gif', '.heic', '.heif']:
            text = self._process_image_convert(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        text = self._clean_text(text)
        
        doc = Document(
            id=str(uuid.uuid4()),
            title=path.stem,
            text=text,
            summary="",
            sections=[]
        )
        
        return doc
    
    def _process_pdf(self, path: Path) -> str:
        """Process PDF - OCR if needed"""
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(str(path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if len(text.strip()) < 50 and PADDLEOCR_AVAILABLE:
                return self._ocr_pdf(path)
            
            return text
        except Exception as e:
            print(f"[DocumentProcessor] PDF text extraction failed: {e}")
            if PADDLEOCR_AVAILABLE:
                return self._ocr_pdf(path)
            raise
    
    def _ocr_pdf(self, path: Path) -> str:
        """Run OCR on PDF"""
        if self.ocr is None:
            print("[DocumentProcessor] Initializing PaddleOCR...")
            self.ocr = PaddleOCR(use_angle_cls=True, lang='en')
        
        from PIL import Image
        import pdf2image
        import numpy as np
        
        images = pdf2image.convert_from_path(str(path))
        text = ""
        
        for i, img in enumerate(images):
            img_array = np.array(img)
            result = self.ocr.ocr(img_array)
            
            if result is None or len(result) == 0:
                continue
            
            ocr_result = result[0]
            if isinstance(ocr_result, dict) and 'rec_texts' in ocr_result:
                texts = ocr_result['rec_texts']
                if texts:
                    text += '\n'.join(texts) + '\n'
                continue
            
            if not ocr_result:
                continue
            
            for line in ocr_result:
                if line:
                    text += line[1][0] + "\n"
        
        return text
    
    def _process_text(self, path: Path) -> str:
        """Process text file"""
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _process_docx(self, path: Path) -> str:
        """Process Word document"""
        try:
            from docx import Document
            doc = Document(str(path.absolute()))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    if any(row_text):
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    def _process_image(self, path: Path) -> str:
        """Process image with OCR - for PNG, JPG, JPEG"""
        if not PADDLEOCR_AVAILABLE:
            raise RuntimeError("PaddleOCR not available. Install: pip install paddleocr")
        
        if self.ocr is None:
            self.ocr = PaddleOCR(use_angle_cls=True, lang='en')
        
        from PIL import Image
        import numpy as np
        img = Image.open(str(path.absolute()))
        img_array = np.array(img)
        
        result = self.ocr.ocr(img_array)
        
        if result is None or len(result) == 0:
            return ""
        
        ocr_result = result[0]
        if isinstance(ocr_result, dict) and 'rec_texts' in ocr_result:
            texts = ocr_result['rec_texts']
            return '\n'.join(texts) if texts else ""
        
        if not ocr_result:
            return ""
        
        text_parts = []
        for line_result in ocr_result:
            if line_result:
                text_parts.append(line_result[1][0])
        
        return '\n'.join(text_parts)
    
    def _process_image_convert(self, path: Path) -> str:
        """Process image with conversion - for GIF, HEIC, HEIF"""
        if not PADDLEOCR_AVAILABLE:
            raise RuntimeError("PaddleOCR not available. Install: pip install paddleocr")
        
        from PIL import Image
        import io
        
        # Open and convert to RGB
        img = Image.open(str(path.absolute()))
        
        # Handle animated GIFs - take first frame
        if img.format == 'GIF' and hasattr(img, 'n_frames'):
            img = img.convert('RGB')
        
        # Convert HEIC/HEIF to PNG in memory
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Save to temporary buffer as PNG
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        
        # Process the converted image using PaddleOCR
        if self.ocr is None:
            self.ocr = PaddleOCR(use_angle_cls=True, lang='en')
        
        import numpy as np
        img_array = np.array(img)
        result = self.ocr.ocr(img_array)
        
        if result is None or len(result) == 0:
            return ""
        
        ocr_result = result[0]
        if isinstance(ocr_result, dict) and 'rec_texts' in ocr_result:
            texts = ocr_result['rec_texts']
            return '\n'.join(texts) if texts else ""
        
        if not ocr_result:
            return ""
        
        text_parts = []
        for line_result in ocr_result:
            if line_result:
                text_parts.append(line_result[1][0])
        
        return '\n'.join(text_parts)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if FTFY_AVAILABLE:
            text = ftfy.fix_text(text)
        
        import re
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into chunks"""
        chunk_size = chunk_size or config.CHUNK_SIZE
        overlap = overlap or config.CHUNK_OVERLAP
        
        chars = list(text)
        chunks = []
        
        for i in range(0, len(chars), chunk_size - overlap):
            chunk = ''.join(chars[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks


class VectorStore:
    """Manages vector storage in Qdrant"""
    
    def __init__(self):
        if not QDRANT_AVAILABLE:
            raise RuntimeError("qdrant-client not available. Install: pip install qdrant-client")
        
        self.client = QdrantClient(path=str(config.VECTOR_DB_DIR))
        self._init_collection()
    
    def _init_collection(self):
        """Initialize or get collection"""
        from qdrant_client.models import VectorParams, Distance
        
        collections = self.client.get_collections().collections
        collection_names = [c.name for c in collections]
        
        if config.COLLECTION_NAME not in collection_names:
            self.client.create_collection(
                collection_name=config.COLLECTION_NAME,
                vectors_config=VectorParams(
                    size=384,
                    distance=Distance.COSINE
                )
            )
    
    def insert_documents(self, doc: Document, doc_embedding: List[float]):
        """Insert document with embedding"""
        points = []
        
        # Document summary
        points.append(PointStruct(
            id=str(uuid.uuid4()),
            vector=doc_embedding,
            payload={
                "type": "document",
                "doc_id": doc.id,
                "title": doc.title,
                "text": doc.summary,
                "doc_title": doc.title,
                "doc_summary": doc.summary,
            }
        ))
        
        # Sections
        for section in doc.sections:
            section_emb = EmbeddingManager.embed_query(section.summary)
            points.append(PointStruct(
                id=section.id,
                vector=section_emb,
                payload={
                    "type": "section",
                    "doc_id": doc.id,
                    "section_id": section.id,
                    "title": section.title,
                    "text": section.text,
                    "summary": section.summary,
                    "doc_title": section.doc_title,
                    "doc_summary": section.doc_summary,
                }
            ))
        
        # Chunks
        if doc.sections:
            for section in doc.sections:
                chunks = DocumentProcessor().chunk_text(section.text)
                for i, chunk in enumerate(chunks):
                    chunk_emb = EmbeddingManager.embed_query(chunk)
                    points.append(PointStruct(
                        id=str(uuid.uuid4()),
                        vector=chunk_emb,
                        payload={
                            "type": "chunk",
                            "doc_id": doc.id,
                            "section_id": section.id,
                            "title": section.title,
                            "text": chunk,
                            "doc_title": section.doc_title,
                            "doc_summary": section.doc_summary,
                        }
                    ))
        
        self.client.upsert(
            collection_name=config.COLLECTION_NAME,
            points=points
        )
    
    def retrieve(self, query_embedding: List[float], top_k: int = 5) -> List[Dict]:
        """Retrieve relevant documents"""
        results = self.client.query_points(
            collection_name=config.COLLECTION_NAME,
            query=query_embedding,
            limit=top_k
        )
        
        return [
            {
                "id": r.id,
                "score": r.score,
                **r.payload
            }
            for r in results.points
        ]


class RAGEngine:
    """Main RAG Engine - ties everything together"""
    
    def __init__(self):
        self.ollama = OllamaManager()
        self.processor = DocumentProcessor()
        self.vector_store = None
    
    def initialize(self):
        """Initialize components"""
        self.ollama.start_service()
        self.vector_store = VectorStore()
    
    def ingest(self, file_path: str) -> Dict[str, Any]:
        """Ingest a document"""
        if self.vector_store is None:
            self.initialize()
        
        doc = self.processor.process_file(file_path)
        
        # Generate document summary using LLM
        summary_prompt = f"Summarize this document in 2-3 sentences:\n\n{doc.text[:1000]}"
        summary = self.ollama.generate(
            summary_prompt,
            system_prompt="You are a helpful assistant that creates concise summaries."
        )
        doc.summary = summary or "Summary not available"
        
        # Create sections (simple splitting by paragraphs)
        paragraphs = doc.text.split('\n\n')
        for i, para in enumerate(paragraphs):
            if len(para.strip()) > 50:
                section = Section(
                    id=str(uuid.uuid4()),
                    title=f"Section {i+1}",
                    text=para.strip(),
                    summary=para.strip()[:200],
                    doc_id=doc.id,
                    doc_title=doc.title,
                    doc_summary=doc.summary
                )
                doc.sections.append(section)
        
        # Generate and store embeddings
        doc_embedding = EmbeddingManager.embed_query(doc.summary)
        self.vector_store.insert_documents(doc, doc_embedding)
        
        # Optionally unload embedding model
        if not config.EMBEDDING_PERSIST:
            EmbeddingManager.unload()
        
        return {
            "status": "success",
            "doc_id": doc.id,
            "title": doc.title,
            "sections": len(doc.sections),
            "summary": doc.summary
        }
    
    def ingest_batch(self, file_paths: List[str]) -> Dict[str, Any]:
        """Ingest multiple files at once
        
        Args:
            file_paths: List of file paths to ingest
            
        Returns:
            Summary of ingestion results
        """
        from pathlib import Path
        
        supported_extensions = {'.pdf', '.txt', '.md', '.docx', '.png', '.jpg', '.jpeg', '.gif', '.heic', '.heif'}
        
        results = []
        errors = []
        processed = 0
        
        for path_str in file_paths:
            path = Path(path_str)
            
            # Check if path exists
            if not path.exists():
                errors.append({
                    "path": path_str,
                    "error": "File not found"
                })
                continue
            
            # If path is a directory, get all supported files
            if path.is_dir():
                for file_path in path.rglob('*'):
                    if file_path.suffix.lower() in supported_extensions:
                        try:
                            result = self.ingest(str(file_path))
                            results.append(result)
                            processed += 1
                        except Exception as e:
                            errors.append({
                                "path": str(file_path),
                                "error": str(e)
                            })
            else:
                # Single file
                if path.suffix.lower() in supported_extensions:
                    try:
                        result = self.ingest(path_str)
                        results.append(result)
                        processed += 1
                    except Exception as e:
                        errors.append({
                            "path": path_str,
                            "error": str(e)
                        })
                else:
                    errors.append({
                        "path": path_str,
                        "error": f"Unsupported file type: {path.suffix}"
                    })
        
        return {
            "status": "success",
            "total_processed": processed,
            "total_errors": len(errors),
            "results": results,
            "errors": errors
        }
    
    def query(self, question: str) -> Dict[str, Any]:
        """Query the RAG system"""
        if self.vector_store is None:
            self.initialize()
        
        # Embed query
        query_embedding = EmbeddingManager.embed_query(question)
        
        # Retrieve relevant documents
        results = self.vector_store.retrieve(
            query_embedding,
            top_k=config.TOP_K_CHUNKS
        )
        
        if not results:
            return {
                "answer": "No relevant documents found. Please ingest some documents first.",
                "sources": []
            }
        
        # Build context from retrieved results - group by document
        docs_context = {}
        for r in results:
            doc_id = r.get('doc_id')
            if doc_id not in docs_context:
                docs_context[doc_id] = {
                    "doc_title": r.get('doc_title', r.get('title', 'Unknown')),
                    "doc_summary": r.get('doc_summary', ''),
                    "chunks": []
                }
            docs_context[doc_id]["chunks"].append({
                "title": r.get('title', ''),
                "text": r.get('text', '')
            })
        
        # Build context with document summaries + relevant chunks
        context_parts = []
        for doc_id, doc_data in docs_context.items():
            context_parts.append(f"Document: {doc_data['doc_title']}")
            if doc_data['doc_summary']:
                context_parts.append(f"Summary: {doc_data['doc_summary']}")
            context_parts.append("Relevant sections:")
            for chunk in doc_data['chunks']:
                context_parts.append(f"  - {chunk['title']}: {chunk['text']}")
            context_parts.append("")
        
        context = "\n\n".join(context_parts)
        
        # Generate answer
        prompt = f"""Based on the following context, answer the question.

Context:
{context}

Question: {question}

Answer:"""
        
        answer = self.ollama.generate(
            prompt,
            system_prompt="You are a helpful assistant. Answer based only on the provided context. If the answer is not in the context, say so."
        )
        
        return {
            "answer": answer,
            "sources": [
                {
                    "doc_title": r.get('doc_title', ''),
                    "doc_summary": r.get('doc_summary', ''),
                    "title": r.get('title', r.get('type', 'Unknown')),
                    "text": r.get('text', '')[:200] + "...",
                    "score": r.get('score', 0)
                }
                for r in results
            ]
        }
    
    def shutdown(self):
        """Shutdown engine and cleanup"""
        self.ollama.stop_service()
        EmbeddingManager.unload()


# Singleton instance
_engine = None

def get_engine() -> RAGEngine:
    """Get or create RAG engine instance"""
    global _engine
    if _engine is None:
        _engine = RAGEngine()
        _engine.initialize()
    return _engine
